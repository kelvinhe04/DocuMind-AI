# -*- coding: utf-8 -*-
"""
Inserta el "Capitulo 20. Demostracion Visual de la Plataforma" en DocuMind_Final.docx,
justo antes de la seccion "Anexos". Usa imagenes reales si existen en docs/captures/<slug>.png,
de lo contrario inserta un recuadro de marcador de posicion.
"""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = r"D:\utp\Innovacion\Semestral\DocuMind AI\docs"
DOC_PATH = os.path.join(DOCS, "DocuMind_Final.docx")
CAP_DIR = os.path.join(DOCS, "captures")

doc = Document(DOC_PATH)


def _ptext(ch):
    return ''.join((t.text or '') for t in ch.findall('.//' + qn('w:t'))).strip()


def remove_existing_chapter():
    """Elimina un 'Capitulo 20' previo (idempotencia) para poder regenerar."""
    body = doc.element.body
    children = list(body)
    start = end = None
    for i, ch in enumerate(children):
        if ch.tag != qn('w:p'):
            continue
        t = _ptext(ch)
        if t.startswith('Capitulo 20. Demostracion Visual') and start is None:
            start = i
        if start is not None and t == 'Anexos':
            end = i
            break
    if start is not None and end is not None:
        for ch in children[start:end]:
            body.remove(ch)
        print("(Capitulo 20 previo eliminado: %d elementos)" % (end - start))


remove_existing_chapter()

# --- localizar la referencia: encabezado H1 "Anexos" ---
ref = None
for p in doc.paragraphs:
    if p.style.name == "Heading 1" and p.text.strip() == "Anexos":
        ref = p
        break
if ref is None:
    raise SystemExit("No se encontro el encabezado 'Anexos'")

GRAY = RGBColor(0x6B, 0x72, 0x80)
fig_counter = {"n": 0}


def _shade_border(p, fill="EEF1F6", color="9CA3AF"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '12')
        e.set(qn('w:space'), '10')
        e.set(qn('w:color'), color)
        pBdr.append(e)
    pPr.append(pBdr)


def heading(text, level):
    p = ref.insert_paragraph_before(text, style="Heading %d" % level)
    return p


def para(text, style="Normal"):
    return ref.insert_paragraph_before(text, style=style)


def bullet(text):
    return ref.insert_paragraph_before(text, style="List Bullet")


def caption(desc):
    fig_counter["n"] += 1
    p = ref.insert_paragraph_before("Figura %d. %s" % (fig_counter["n"], desc), style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def figure(slug, desc):
    """Imagen real si existe; si no, recuadro marcador. Luego el pie de figura."""
    img = os.path.join(CAP_DIR, slug + ".png")
    if not os.path.exists(img):
        img = os.path.join(CAP_DIR, slug + ".jpg")
    p = ref.insert_paragraph_before(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img):
        run = p.add_run()
        run.add_picture(img, width=Cm(15.5))
    else:
        _shade_border(p)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("[ CAPTURA PENDIENTE ]")
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        run.add_break()
        r2 = p.add_run("Reemplazar por: %s" % desc)
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = GRAY
        for _ in range(4):
            r2.add_break()
        r3 = p.add_run("(slug sugerido: captures/%s.png)" % slug)
        r3.italic = True
        r3.font.size = Pt(8)
        r3.font.color.rgb = GRAY
    caption(desc)


def pricing_table():
    rows = [
        ("Plan", "Precio mensual", "Equivalente anual*", "Dirigido a", "Prestaciones principales"),
        ("Free", "$0 / siempre", "$0", "Empezar a explorar",
         "3 documentos, 20 consultas/mes, solo PDF con texto, busqueda semantica, soporte por email"),
        ("Starter (Mas popular)", "$49 / mes", "$41/mes ($488/ano)", "Equipos pequenos",
         "Documentos y consultas ilimitados, OCR (imagenes y PDF escaneados), chat con citas, analytics basicos, soporte prioritario"),
        ("Business", "$199 / mes", "$165/mes ($1,982/ano)", "Empresas medianas",
         "Todo de Starter + multi-workspace (3), roles y permisos avanzados, analytics avanzados, acceso API, SLA 99.9%"),
        ("Enterprise", "$999 / mes", "$829/mes ($9,950/ano)", "Infraestructura propia",
         "Todo de Business + workspaces ilimitados, despliegue on-premise, SSO/SAML, integracion custom, soporte dedicado 24/7"),
    ]
    table = doc.add_table(rows=0, cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, rowdata in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(rowdata):
            cells[ci].text = val
            for pp in cells[ci].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(9)
                    if ri == 0:
                        rr.bold = True
            if ri == 0:
                tcPr = cells[ci]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '7C3AED')
                tcPr.append(shd)
                for pp in cells[ci].paragraphs:
                    for rr in pp.runs:
                        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # mover la tabla a su lugar (antes de la referencia)
    ref._p.addprevious(table._tbl)


# =========================================================================
#  CONTENIDO
# =========================================================================
heading("Capitulo 20. Demostracion Visual de la Plataforma", 1)
para("Este capitulo presenta la evidencia visual del prototipo funcional de DocuMind AI. "
     "Se documentan, mediante capturas de pantalla, tanto la pagina de aterrizaje publica "
     "—incluyendo el detalle de los planes de cobro— como cada una de las secciones del panel "
     "de control disponible para los usuarios registrados. El producto esta desarrollado con "
     "Next.js 14 y React en el frontend, autenticacion con Clerk, cobros con Stripe y un backend "
     "en FastAPI que implementa recuperacion aumentada por generacion (RAG) con busqueda hibrida "
     "(BM25 + vectorial). Las imagenes incluidas a continuacion corresponden a la interfaz real "
     "de la aplicacion.")

# ---------------- LANDING ----------------
heading("Pagina de Aterrizaje (Landing Page)", 2)
para("La pagina de aterrizaje es el punto de entrada publico de DocuMind AI (ruta /). Emplea un "
     "diseno oscuro con acentos violeta y cian; su objetivo es comunicar la propuesta de valor, "
     "generar confianza y conducir al visitante hacia el registro o la contratacion de un plan. "
     "A continuacion se documentan sus secciones en el orden en que aparecen al desplazarse por la pagina.")

heading("Vista general de la landing", 3)
para("Captura completa (full page) de la pagina de aterrizaje, de extremo a extremo, que permite "
     "apreciar la composicion general y la jerarquia visual del sitio.")
figure("landing-full", "Vista completa de la pagina de aterrizaje de DocuMind AI")

heading("Barra de navegacion (Navbar)", 3)
para("Encabezado fijo con el logotipo de DocuMind AI y accesos directos a Caracteristicas, Como "
     "funciona y Precios, junto con los botones de Iniciar sesion y Registrarse. Permanece visible "
     "durante todo el desplazamiento para facilitar la navegacion.")
figure("landing-navbar", "Barra de navegacion superior de la landing")

heading("Seccion principal (Hero)", 3)
para("Encabezado principal con el mensaje de valor central, un subtitulo explicativo y los botones "
     "de accion primarios. Incluye una demostracion visual del chat documental respondiendo "
     "preguntas como «Dias de vacaciones», «Canon de arrendamiento» o «Causales de despido» con "
     "citas a la fuente.")
figure("landing-hero", "Seccion Hero con la propuesta de valor y demo del chat")

heading("Caracteristicas (Features)", 3)
para("Cuadricula de funcionalidades clave: busqueda semantica, chat con citas exactas, OCR "
     "inteligente, analytics en tiempo real, multi-workspace y despliegue on-premise. Cada tarjeta "
     "resume el beneficio y, cuando corresponde, muestra una etiqueta de «Proximamente».")
figure("landing-features", "Cuadricula de caracteristicas del producto")

heading("Como funciona (HowItWorks)", 3)
para("Explicacion del flujo en tres pasos: (1) Sube tus documentos, (2) Pregunta en lenguaje "
     "natural y (3) Obten respuestas con citas. Comunica la simplicidad de uso del producto.")
figure("landing-howitworks", "Flujo de uso en tres pasos")

heading("Video demostrativo (DemoVideo)", 3)
para("Reproductor de la demostracion funcional que recorre la subida con OCR, el chat con citas, "
     "la busqueda semantica y el pago con Stripe en vivo.")
figure("landing-demo", "Bloque de video demostrativo")

heading("Metricas e indicadores (Stats)", 3)
para("Bloque de cifras destacadas: documentos procesados, preguntas respondidas, empresas en "
     "lista de espera y precision en las citas. Aporta prueba social cuantitativa.")
figure("landing-stats", "Indicadores y metricas destacadas")

heading("Planes y precios (Pricing)", 3)
para("DocuMind AI emplea un modelo de suscripcion SaaS por niveles. En la landing, los planes se "
     "presentan con un conmutador entre facturacion mensual y anual; la facturacion anual aplica "
     "un descuento del 17 %. La siguiente tabla resume los cuatro planes, su precio y las "
     "prestaciones incluidas en cada uno.")
pricing_table()
para("*El equivalente anual corresponde al precio mensual con el descuento del 17 % aplicado; "
     "entre parentesis se indica el total facturado por ano. Nota: todos los planes incluyen el "
     "modo extractivo como respaldo (fallback), por lo que la aplicacion sigue respondiendo aun "
     "sin la clave de API de Groq.")
figure("landing-pricing", "Seccion de planes y precios de la landing")

heading("Testimonios (Testimonials)", 3)
para("Opiniones de perfiles representativos —Gerente de RR. HH., Socia Senior de un bufete y "
     "Director de Biblioteca Digital— que ilustran casos de uso reales en los segmentos objetivo.")
figure("landing-testimonials", "Testimonios de clientes")

heading("Llamado a la accion (CTA)", 3)
para("Seccion de cierre que invita a comenzar gratis o a contactar al equipo de ventas, con un "
     "boton de accion destacado.")
figure("landing-cta", "Seccion de llamado a la accion (CTA)")

heading("Pie de pagina (Footer)", 3)
para("Pie con enlaces organizados por Producto y Empresa (Caracteristicas, Precios, Roadmap, "
     "Acerca de, Contacto, Privacidad, Terminos) y datos de la marca.")
figure("landing-footer", "Pie de pagina de la landing")

# ---------------- AUTENTICACION ----------------
heading("Registro e Inicio de Sesion (Clerk)", 2)
para("El acceso al panel de control esta protegido mediante Clerk. Los visitantes deben crear una "
     "cuenta o iniciar sesion antes de utilizar las funcionalidades de la aplicacion.")

heading("Registro de cuenta (Sign up)", 3)
para("Formulario de creacion de cuenta (ruta /sign-up) gestionado por Clerk, con registro por "
     "correo electronico y proveedores de autenticacion social.")
figure("auth-signup", "Pantalla de registro de cuenta")

heading("Inicio de sesion (Sign in)", 3)
para("Pantalla de autenticacion (ruta /sign-in) para usuarios existentes.")
figure("auth-signin", "Pantalla de inicio de sesion")

# ---------------- DASHBOARD ----------------
heading("Panel de Control (Dashboard)", 2)
para("Tras autenticarse, el usuario accede al area privada de la aplicacion. La interfaz se "
     "organiza con una barra lateral de navegacion (Dashboard, Chat, Documentos, Subir, Busqueda "
     "y Facturacion) y una barra superior con el titulo de la seccion, el indicador del plan "
     "vigente y el menu de usuario. A continuacion se documenta cada pantalla.")

heading("Dashboard y panel de metricas", 3)
para("Vista principal (ruta /dashboard) con indicadores de uso —documentos, fragmentos indexados, "
     "consultas realizadas y almacenamiento— frente a los limites del plan, junto con accesos "
     "rapidos para subir un documento, abrir el chat, buscar o auditar la base. Incluye un grafico "
     "de actividad reciente.")
figure("app-dashboard", "Panel de metricas y accesos rapidos")

heading("Chat documental", 3)
para("Conversacion en lenguaje natural sobre los documentos cargados (ruta /chat). El sistema "
     "responde mediante RAG (Groq) y muestra citas a las fuentes; ofrece sugerencias como «Resume "
     "el ultimo contrato cargado» o «Encuentra riesgos o clausulas importantes», historial de "
     "conversaciones y la opcion de compartir el chat mediante un enlace.")
figure("app-chat", "Interfaz de chat documental con citas a las fuentes")

heading("Documentos", 3)
para("Listado de los archivos cargados (ruta /documents) con su numero de paginas y fragmentos, "
     "buscador por nombre, vista previa y opciones de eliminacion individual y masiva.")
figure("app-documents", "Listado y gestion de documentos")

heading("Subir documento", 3)
para("Zona de carga por arrastrar y soltar (ruta /upload) para archivos PDF, JPG o PNG, que "
     "muestra el progreso del procesamiento: en cola, extraccion de texto (OCR cuando corresponde), "
     "preparacion de vectores e indexacion del contenido.")
figure("app-upload", "Zona de carga de documentos con progreso de procesamiento")

heading("Busqueda semantica", 3)
para("Buscador (ruta /search) que recupera los fragmentos mas relevantes de la base documental "
     "mediante busqueda hibrida (BM25 + vectorial), mostrando el documento de origen y el pasaje "
     "correspondiente.")
figure("app-search", "Resultados de la busqueda semantica")

heading("Facturacion", 3)
para("Resumen del plan contratado (ruta /billing), consumo frente a los limites (documentos, "
     "consultas y almacenamiento), historial y gestion de la suscripcion mediante Stripe.")
figure("app-billing", "Panel de facturacion y consumo del plan")

heading("Planes desde la aplicacion", 3)
para("Vista interna de planes (ruta /pricing) que permite al usuario autenticado comparar niveles "
     "y mejorar su suscripcion; la contratacion se procesa con Stripe Checkout.")
figure("app-pricing", "Comparativa de planes dentro de la aplicacion")

heading("Pago y confirmacion (Stripe)", 3)
para("Pantalla de Stripe Checkout para el cobro y pagina de confirmacion mostrada tras un pago "
     "exitoso, que actualiza automaticamente el plan del usuario.")
figure("app-stripe", "Proceso de pago con Stripe Checkout y confirmacion")

doc.save(DOC_PATH)
print("OK -> %s   |   figuras insertadas: %d" % (DOC_PATH, fig_counter["n"]))
