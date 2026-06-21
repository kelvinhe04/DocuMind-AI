"""
Genera DocuMind_Presentacion_Semestral.docx

Uso:
  cd "DocuMind AI"
  python docs/generate_word.py

Requiere: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = os.path.join(os.path.dirname(__file__), "DocuMind_Presentacion_Semestral.docx")

# ── Colores de marca ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x1B, 0x4B)   # Azul marino UTP
BLUE   = RGBColor(0x00, 0x70, 0xC0)   # Azul Microsoft — subtítulos nivel 2
VIOLET = RGBColor(0x5B, 0x21, 0xB6)   # Violet — citas/acento en cuerpo
DARK   = RGBColor(0x1E, 0x1E, 0x2E)   # Texto oscuro
GRAY   = RGBColor(0x64, 0x64, 0x80)   # Gris secundario
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GOLD   = RGBColor(0xD9, 0xA8, 0x00)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _paragraph_spacing(para, before=0, after=6, line=None) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_toc(doc: Document) -> None:
    """Inserta un campo TOC que Word actualiza al abrir."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    # Placeholder text
    run2 = para.add_run("[ Actualice el índice: clic derecho → Actualizar campo ]")
    run2.font.color.rgb = GRAY
    run2.font.italic = True
    run2.font.size = Pt(10)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = BLUE
        run.font.size = Pt(13)
        run.font.bold = True
    else:
        run.font.color.rgb = DARK
        run.font.size = Pt(11)
        run.font.bold = True
    _paragraph_spacing(h, before=120, after=60)


def _body(doc: Document, text: str, bold=False, italic=False,
          color=None, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    _paragraph_spacing(p, before=0, after=60)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    _paragraph_spacing(p, before=0, after=30)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)


def _table(doc: Document, headers: list, rows: list[list],
           header_bg="1E1E2E", header_color=WHITE,
           col_widths: list = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabecera
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _set_cell_bg(cell, header_bg)
        _set_cell_borders(cell, "444444")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = header_color

    # Filas de datos
    for r_idx, row in enumerate(rows):
        bg = "F8F8FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, "CCCCCC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)

    # Ancho de columnas
    if col_widths:
        for r_idx, row in enumerate(t.rows):
            for c_idx, cell in enumerate(row.cells):
                if c_idx < len(col_widths):
                    cell.width = Cm(col_widths[c_idx])

    doc.add_paragraph()


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

def build_document() -> None:
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ─────────────────────────────────────────────────────────────────────────
    # 0. PORTADA
    # ─────────────────────────────────────────────────────────────────────────
    def _cen(text, bold=False, size=12, color=None, space_after=120):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _paragraph_spacing(p, before=0, after=space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    # Espacio superior
    for _ in range(2):
        _cen("", size=12)

    _cen("UNIVERSIDAD TECNOLÓGICA DE PANAMÁ", bold=True, size=14, color=NAVY, space_after=60)
    _cen("FACULTAD DE INGENIERÍA DE SISTEMAS COMPUTACIONALES", bold=True, size=12, color=NAVY, space_after=60)
    _cen("DESARROLLO Y GESTIÓN DE SOFTWARE", bold=False, size=11, color=DARK, space_after=60)
    _cen("INNOVACIÓN Y EMPRENDIMIENTO", bold=False, size=11, color=DARK, space_after=120)

    _cen("CASO DE INNOVACIÓN", bold=True, size=16, color=BLUE, space_after=200)

    # Título del proyecto
    _cen("DocuMind AI", bold=True, size=26, color=NAVY, space_after=60)
    _cen('"El cerebro digital de tu empresa"', bold=False, size=13, color=GRAY, space_after=200)

    _cen("Grupo: 1GS241", bold=True, size=12, color=DARK, space_after=60)

    _cen("Estudiantes:", bold=True, size=12, color=DARK, space_after=40)
    for nombre_cedula in [
        "He, Kelvin — 8-999-1950",
        "Barrera, Roy — 8-1022-2121",
        "Mosquera, Einer — 8-924-1880",
        "Athanasidis, Nicolás — 8-1001-974",
    ]:
        _cen(nombre_cedula, bold=False, size=11, color=DARK, space_after=30)

    _cen("", size=11, space_after=60)
    _cen("Profe: Melvin Falcón", bold=False, size=11, color=DARK, space_after=30)
    _cen("Fecha: 29/5/2026", bold=False, size=11, color=DARK, space_after=30)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 1. ÍNDICE
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "Índice", level=1)
    _add_toc(doc)
    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 2. QUIÉNES SOMOS — LA STARTUP
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "1. Quiénes Somos — La Startup", level=1)

    _heading(doc, "1.1 Nombre y Eslogan", level=2)
    _body(doc, "Nombre: DocuMind AI", bold=True)
    _body(doc, 'Eslogan: "El cerebro digital de tu empresa."', italic=True)

    _heading(doc, "1.2 Visión", level=2)
    _body(doc,
          "Convertirnos en la infraestructura de conocimiento estándar para empresas "
          "latinoamericanas, comenzando por Panamá y expandiéndonos por Centroamérica y el Caribe.")

    _heading(doc, "1.3 Misión", level=2)
    _body(doc,
          "Democratizar el acceso a la inteligencia documental para empresas de todos los tamaños, "
          "eliminando la pérdida de tiempo en la búsqueda de información y reduciendo errores por "
          "uso de documentos obsoletos.")

    # ─────────────────────────────────────────────────────────────────────────
    # 3. ORGANIGRAMA
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "2. Organigrama del Equipo", level=1)

    _body(doc,
          "El equipo de DocuMind AI está conformado por cuatro integrantes del grupo 1GS241, "
          "cada uno con un rol claramente definido:")

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _paragraph_spacing(p_org, before=60, after=60)
    run_org = p_org.add_run(
        "                     ┌─────────────────────┐\n"
        "                     │        CEO          │\n"
        "                     │    Kelvin He        │\n"
        "                     │ (Visión, Estrategia,│\n"
        "                     │  Pitch, Relaciones) │\n"
        "                     └──────────┬──────────┘\n"
        "                                │\n"
        "          ┌─────────────────────┼─────────────────────┐\n"
        "          │                     │                     │\n"
        "  ┌───────┴───────┐    ┌────────┴────────┐    ┌───────┴───────┐\n"
        "  │   Backend     │    │    Frontend     │    │      QA       │\n"
        "  │  Roy Barrera  │    │ Einer Mosquera  │    │  Nicolás      │\n"
        "  │ (FastAPI, RAG,│    │(Next.js, React, │    │  Athanasidis  │\n"
        "  │  ChromaDB,    │    │ shadcn/ui,      │    │(Testing, Docs,│\n"
        "  │  Embeddings)  │    │ Tailwind, Demo) │    │ Pitch Support)│\n"
        "  └───────────────┘    └─────────────────┘    └───────────────┘"
    )
    run_org.font.name = "Courier New"
    run_org.font.size = Pt(8)
    run_org.font.color.rgb = DARK

    _heading(doc, "Roles Detallados", level=2)
    _table(doc,
           ["Rol", "Integrante", "Responsabilidades"],
           [
               ["CEO", "Kelvin He",
                "Visión del producto, pitch a inversionistas, estrategia de negocio, "
                "relaciones con aliados (AMPYME, Cámara de Comercio), toma de decisiones ejecutivas."],
               ["Backend Lead", "Roy Barrera",
                "Arquitectura del pipeline RAG, API REST con FastAPI, ChromaDB, embeddings, "
                "integración con LLMs (Groq), seguridad de datos."],
               ["Frontend Lead", "Einer Mosquera",
                "Interfaz de usuario con Next.js y shadcn/ui, experiencia de chat, dashboard de "
                "métricas, responsive design, animaciones y polish visual."],
               ["QA / Documentación", "Nicolás Athanasidis",
                "Pruebas funcionales de la demo, documentación técnica, soporte en el pitch, "
                "control de calidad del producto, manuales de usuario."],
           ],
           col_widths=[3.0, 3.5, 9.5])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 4. PROBLEMA
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "3. Problema que Resolvemos", level=1)

    _heading(doc, "3.1 Descripción del Problema", level=2)
    _body(doc,
          "Las empresas en Panamá y Latinoamérica almacenan miles de documentos digitales "
          "sin una forma inteligente de consultarlos: contratos laborales, manuales de procedimientos, "
          "normativas legales, facturas y correspondencia institucional.")
    _body(doc, "Cuando un empleado necesita información:")
    for item in [
        "Pierde tiempo buscando en carpetas compartidas desorganizadas.",
        "Pregunta a otros compañeros, interrumpiendo su productividad.",
        "Puede utilizar una versión incorrecta u obsoleta del documento.",
        "No hay forma de verificar si la información encontrada es la más reciente.",
    ]:
        _bullet(doc, item)

    _heading(doc, "3.2 Datos de Impacto", level=2)
    _body(doc, "En una empresa de 50 colaboradores:", bold=True)
    _table(doc,
           ["Métrica", "Valor"],
           [
               ["Horas perdidas por empleado buscando docs/semana", "2.5 horas"],
               ["Total horas semanales desperdiciadas", "125 horas"],
               ["Horas anuales de productividad perdida", "+6,000 horas"],
               ["Costo estimado anual en tiempo no productivo", "$45,000 – $75,000 USD"],
           ],
           col_widths=[11.0, 5.0])

    _heading(doc, "3.3 Contexto Panameño", level=2)
    for item in [
        "El 95% de las empresas en Panamá son MIPYMES sin herramientas tecnológicas para gestionar su conocimiento documental. (AMPYME, 2026)",
        "Las MIPYMES representan más del 90% de las empresas en Panamá y el 70% de la economía. (La Prensa Panamá, 2026)",
        "El 95% de las empresas mueren en sus primeros tres años por falta de preparación y procesos, no por falta de capital.",
    ]:
        _bullet(doc, item)

    _heading(doc, "3.4 Casos Reales en Panamá", level=2)
    for item in [
        "Un bufete de abogados en Panamá City tiene 15,000 resoluciones de la CSJ en PDFs sin índice.",
        "Una constructora local pierde 3 horas semanales por empleado buscando cláusulas de contratos.",
        "Una universidad privada no puede responder rápidamente preguntas sobre reglamentos dispersos en cientos de PDFs.",
    ]:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 5. EL SOFTWARE
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "4. El Software — DocuMind AI", level=1)

    _heading(doc, "4.1 Descripción del Producto", level=2)
    _body(doc,
          "DocuMind AI es un Sistema Operativo del Conocimiento Empresarial que convierte "
          "documentos estáticos en conocimiento accionable mediante inteligencia artificial. "
          "No es un chatbot genérico: responde con los documentos internos de tu empresa, "
          "citando archivo, página y párrafo exactos.")

    _heading(doc, "4.2 Funcionalidades Principales", level=2)
    for item in [
        ("DocuMind Search: ", "Búsqueda semántica que entiende el significado, no solo palabras clave."),
        ("DocuMind Chat: ", "Chatbot conversacional que responde en lenguaje natural citando la fuente exacta."),
        ("DocuMind Analytics: ", "Dashboard empresarial con métricas de uso, productividad y tendencias."),
        ("DocuMind Insights: ", "(Roadmap) Análisis automático para detectar riesgos y vencimientos."),
        ("DocuMind Compliance: ", "(Roadmap) Monitoreo de normativas con alertas automáticas."),
    ]:
        _bullet(doc, item[1], bold_prefix=item[0])

    _heading(doc, "4.3 Diferenciación Clave", level=2)
    _body(doc,
          '"ChatGPT responde con conocimiento general de internet. DocuMind responde con los '
          'documentos internos de tu empresa, citando archivo, página y párrafo exactos."',
          italic=True, color=VIOLET)

    _heading(doc, "4.4 Tecnologías Utilizadas", level=2)
    _table(doc,
           ["Capa", "Tecnología", "Versión", "Propósito"],
           [
               ["Frontend", "Next.js + React + TypeScript", "14.2 / 18", "App Router, SSR, páginas de la app"],
               ["Frontend", "Tailwind CSS v4 + shadcn/ui", "4.3 / 4.11", "Estilos utilitarios + componentes UI"],
               ["Frontend", "Framer Motion", "12", "Animaciones de la landing y app"],
               ["Frontend", "Recharts", "3", "Gráficas del dashboard de métricas"],
               ["Auth", "Clerk", "7", "Autenticación real (OAuth + email)"],
               ["Pagos", "Stripe", "22", "Suscripciones en modo test"],
               ["Backend", "FastAPI + Uvicorn", "0.138", "4 endpoints REST con hot reload"],
               ["Backend", "Python", "3.11", "Lenguaje del backend"],
               ["IA – Embeddings", "all-MiniLM-L6-v2 (ONNX)", "22 MB", "Embeddings locales sin PyTorch"],
               ["IA – LLM", "Groq (llama-3.1-8b-instant)", "gratis", "30 req/min, ~1s respuesta"],
               ["IA – Fallback", "Modo extractivo", "—", "Sin LLM: muestra chunks relevantes"],
               ["Vector DB", "ChromaDB", "1.5", "Búsqueda vectorial persistente local"],
               ["Metadata DB", "SQLite", "3.x", "Metadatos de documentos y consultas"],
               ["OCR", "Tesseract + pdf2image + Pillow", "—", "PDFs escaneados e imágenes JPG/PNG"],
           ],
           col_widths=[3.0, 5.5, 2.5, 5.0])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 6. LANDING PAGE
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "5. Landing Page", level=1)

    _heading(doc, "5.1 Descripción", level=2)
    _body(doc,
          "Landing page moderna y futurista (estilo SaaS: Vercel, Linear, Notion) integrada "
          "en el mismo proyecto Next.js, accesible en la ruta raíz /. Sirve como cara pública "
          "de DocuMind AI y canal principal de adquisición.")

    _heading(doc, "5.2 Secciones de la Landing", level=2)
    for item in [
        ("Navbar: ", "Transparente → sólida al hacer scroll. Logo, links, CTA 'Iniciar sesión'."),
        ("Hero: ", "Título animado, eslogan, badge 'Ahora con OCR', CTA → /sign-up, screenshot flotante del producto."),
        ("Logo Cloud: ", "Logos de tecnologías utilizadas (Groq, Clerk, Stripe, ChromaDB, Tesseract)."),
        ("Features: ", "6 tarjetas animadas: búsqueda semántica, chat con citas, OCR, analytics, multi-workspace, on-premise."),
        ("HowItWorks: ", "3 pasos ilustrados: Sube, Pregunta, Obtén la respuesta con fuente exacta."),
        ("Pricing: ", "4 planes (Free, Starter $49, Business $199, Enterprise $999), toggle mensual/anual."),
        ("Testimonials: ", "3 testimonios de empresas panameñas ficticias (Nova S.A., Bufete Méndez, U. del Pacífico)."),
        ("Stats: ", '"500+ docs procesados", "10k+ preguntas respondidas", "50+ en waitlist".'),
        ("CTA: ", "Formulario de email para waitlist: 'Acceso anticipado gratuito para las primeras 50 empresas'."),
        ("Footer: ", "Links, redes sociales, política de privacidad."),
    ]:
        _bullet(doc, item[1], bold_prefix=item[0])

    _heading(doc, "5.3 Diseño Visual", level=2)
    for item in [
        "Fondo: slate-950 → slate-900 con grid sutil de puntos.",
        "Acentos: violet-500 (primario), emerald-500 (CTA), amber-500 (badge OCR).",
        "Tipografía: Inter (sans-serif moderna, libre).",
        "Animaciones: Framer Motion — fade-in al scroll, slide-up en cards, hover scale+glow.",
        "Efectos: glassmorphism en navbar/cards, gradient border en plan recomendado.",
    ]:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 7. DEMOSTRACIÓN
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "6. Demostración", level=1)

    _heading(doc, "6.1 Escenario", level=2)
    _body(doc,
          "Constructora Nova S.A. — empresa ficticia de construcción en Panamá con 50 empleados. "
          "Maneja contratos laborales, manuales de RRHH, reglamentos internos, contratos de "
          "arrendamiento, políticas de compras y actas de reunión.")

    _heading(doc, "6.2 Documentos Precargados (7 documentos)", level=2)
    _table(doc,
           ["ID", "Archivo", "Tipo", "Contenido Clave"],
           [
               ["001", "Contrato_Laboral_Nova_2026.pdf", "PDF texto",
                "Preaviso 30 días (p.3), seguro médico emergencias 100% (p.4)"],
               ["002", "Manual_RRHH_Nova.pdf", "PDF texto",
                "Horario L-V 8am-5pm (p.2), 17 días vacaciones a 2 años (p.4), causales despido (p.6)"],
               ["003", "Reglamento_Interno_Nova.pdf", "PDF texto",
                "Renuncia sin preaviso → pérdida de beneficios (p.3)"],
               ["004", "Contrato_Arrendamiento_Oficina.pdf", "PDF texto",
                "Plazo 24 meses renovables (p.2), canon $3,500 USD/mes (p.3)"],
               ["005", "Politica_Compras_Proveedores.pdf", "PDF texto",
                "3 cotizaciones para compras >$500 (p.3), aprueba GG+Finanzas para $3,000 (p.4)"],
               ["006", "Acta_Reunion_Directiva_Nova.jpg", "OCR imagen",
                "Presupuesto anual 2026 aprobado: $3,200,000 (procesado con OCR)"],
               ["007", "Factura_Servicios_Enero_2026.pdf", "OCR PDF",
                "Total a pagar: $8,160.00 USD (factura escaneada, procesada con OCR)"],
           ],
           col_widths=[1.0, 5.0, 2.5, 7.5])

    _heading(doc, "6.3 Flujo de Demo — 3 Minutos", level=2)
    steps = [
        ("Paso 1 — Landing (15 s): ",
         "Abrir http://localhost:3000. Mostrar landing futurista animada."),
        ("Paso 2 — Login Clerk (10 s): ",
         "Clic en 'Iniciar sesión'. Registrarse con email real. Redirige al dashboard."),
        ("Paso 3 — Dashboard (20 s): ",
         "7 documentos indexados, métricas reales, gráfica de actividad."),
        ("Paso 4 — OCR en vivo (30 s): ",
         "Subir Acta_Reunion_Directiva_Nova.jpg. Ver progreso: 'Procesando imagen… → OCR… → Embeddings…'. "
         "Resultado: 1 página, 12 chunks, confianza OCR 90%."),
        ("Paso 5 — Chat con citas (90 s): ",
         "Preguntas: '¿Vacaciones con 2 años?' → 17 días hábiles. "
         "'¿Qué dice el acta sobre el presupuesto?' → respuesta OCR con nota. "
         "'¿Plazo del arrendamiento?' → 24 meses."),
        ("Paso 6 — Búsqueda semántica (20 s): ",
         "Buscar 'procedimiento disciplinario'. 4 resultados con score de relevancia."),
        ("Paso 7 — Free vs Pago (15 s): ",
         "Plan Free topa 3 documentos → UpgradeDialog → Stripe Checkout → pago con 4242… → "
         "plan sube a Starter en vivo sin ngrok."),
        ("Paso 8 — Cierre (10 s): ",
         "7 documentos procesados, incluyendo OCR. Imaginen miles de documentos y decenas de "
         "empleados consultando simultáneamente."),
    ]
    for label, desc in steps:
        _bullet(doc, desc, bold_prefix=label)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 8. CÓMO FUNCIONA EL SOFTWARE
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "7. Cómo Funciona el Software", level=1)

    _heading(doc, "7.1 Los 4 Endpoints de FastAPI", level=2)
    _table(doc,
           ["#", "Método", "Ruta", "Función"],
           [
               ["1", "POST", "/documents/upload",
                "Extrae texto (o aplica OCR), genera chunks, embeddings e indexa en ChromaDB"],
               ["2", "GET", "/documents",
                "Lista todos los documentos con metadata (tipo, páginas, chunks, fecha)"],
               ["3", "POST", "/chat",
                "RAG completo: mode=chat usa Groq; mode=search devuelve chunks; sin Groq → extractivo"],
               ["4", "GET", "/metrics",
                "Métricas del dashboard: documentos, chunks, consultas, storage, actividad 7 días"],
           ],
           col_widths=[0.8, 2.0, 5.0, 8.2])

    _heading(doc, "7.2 Pipeline RAG (Retrieval-Augmented Generation)", level=2)
    _body(doc, "Flujo de ingesta (POST /documents/upload):", bold=True)
    for step in [
        "Recibir archivo (PDF / JPG / PNG).",
        "Detectar tipo: PDF-texto (pymupdf) vs. PDF-escaneado / imagen (OCR).",
        "Si OCR: preprocesar imagen (escala de grises, binarización) → Tesseract extrae texto.",
        "Limpiar texto y dividir en chunks de 1000 tokens con solapamiento de 200 (RecursiveCharacterTextSplitter).",
        "Generar embeddings con all-MiniLM-L6-v2 vía ChromaDB ONNX (sin PyTorch, ~22 MB).",
        "Indexar chunks en ChromaDB persistente con metadata {filename, page, chunk_index}.",
        "Registrar documento en SQLite.",
    ]:
        _bullet(doc, step)

    _body(doc, "Flujo de consulta (POST /chat mode=chat):", bold=True)
    for step in [
        "Recibir pregunta en lenguaje natural.",
        "Convertir pregunta a embedding y buscar los top-5 chunks más similares en ChromaDB.",
        "Construir prompt con contexto (chunks) + pregunta y enviarlo a Groq (llama-3.1-8b-instant).",
        "Groq genera respuesta con cita: 'Fuente: [archivo], página X.'",
        "Si Groq no está disponible → modo extractivo: devuelve los chunks directamente (la app nunca se cae).",
        "Registrar consulta en SQLite para métricas.",
    ]:
        _bullet(doc, step)

    _heading(doc, "7.3 Arquitectura de Seguridad y Gating de Planes", level=2)
    _body(doc,
          "El gating Free vs. Pago se implementa en la capa proxy de Next.js "
          "(src/app/api/proxy/[...path]/route.ts), no en FastAPI. Esto mantiene FastAPI 'limpio' "
          "con exactamente 4 endpoints:")
    for item in [
        "Plan Free: máximo 3 documentos, sin OCR (imágenes bloqueadas).",
        "Plan Starter ($49/mes): documentos ilimitados, OCR activo.",
        "Plan Business ($199/mes) y Enterprise ($999/mes): features avanzados.",
        "El plan se almacena en publicMetadata de Clerk → sin base de datos propia de planes.",
        "Al exceder límites: el proxy devuelve HTTP 402 → el frontend abre UpgradeDialog → Stripe Checkout.",
    ]:
        _bullet(doc, item)

    _heading(doc, "7.4 Modelos de IA Utilizados", level=2)
    _table(doc,
           ["Modelo", "Proveedor", "Uso", "Costo"],
           [
               ["llama-3.1-8b-instant", "Groq API", "Generación de respuestas en lenguaje natural", "Gratis (30 req/min)"],
               ["all-MiniLM-L6-v2", "ONNX / ChromaDB", "Embeddings para búsqueda semántica", "Gratis (local, 22 MB)"],
               ["Modo Extractivo", "Local", "Fallback sin LLM: devuelve chunks directos", "Sin costo"],
           ],
           col_widths=[4.5, 3.5, 5.5, 2.5])

    _heading(doc, "7.5 OCR y Procesamiento de Imágenes", level=2)
    _body(doc,
          "DocuMind AI incluye soporte completo para documentos escaneados e imágenes mediante "
          "Tesseract OCR, el motor de reconocimiento óptico de caracteres open-source más utilizado del mundo.")
    _table(doc,
           ["Tipo de Archivo", "Procesamiento", "OCR Requerido"],
           [
               ["PDF con texto plano (.pdf)", "pymupdf extrae texto directamente", "NO"],
               ["PDF escaneado (.pdf)", "pdf2image convierte páginas a imágenes → Tesseract OCR", "SÍ"],
               ["Imagen (.jpg, .jpeg, .png)", "Preprocesamiento Pillow → Tesseract OCR", "SÍ"],
           ],
           col_widths=[5.5, 7.5, 3.0])
    _body(doc, "Pipeline de preprocesamiento OCR:", bold=True)
    for step in [
        "Convertir imagen a escala de grises (grayscale).",
        "Binarización simple: píxeles < 140 → negro, resto → blanco (mejora contraste).",
        "Ejecutar Tesseract en idioma español (spa) con fallback a inglés (eng).",
        "Calcular confianza promedio del OCR y almacenarla en la metadata del documento.",
    ]:
        _bullet(doc, step)

    _heading(doc, "7.6 Autenticación y Pagos (Clerk + Stripe)", level=2)
    _body(doc, "Autenticación con Clerk:", bold=True)
    for item in [
        "Clerk maneja registro, login y sesiones (OAuth Google + email/contraseña).",
        "Middleware de Next.js protege rutas /dashboard, /chat, /documents, /upload, /search, /billing.",
        "El plan de suscripción se almacena en publicMetadata.plan de Clerk (sin BD adicional).",
    ]:
        _bullet(doc, item)

    _body(doc, "Pagos con Stripe:", bold=True)
    for item in [
        "Stripe Checkout crea sesiones de pago en modo test (sin webhooks ni ngrok).",
        "Al completar el pago, /api/stripe/confirm verifica la sesión y actualiza el plan en Clerk.",
        "El upgrade Free → Starter es instantáneo y visible en la UI sin recargar manualmente.",
        "Tarjeta de prueba: 4242 4242 4242 4242, fecha futura, CVC cualquiera.",
    ]:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 9. MERCADO OBJETIVO
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "8. Mercado Objetivo", level=1)

    _heading(doc, "8.1 TAM / SAM / SOM", level=2)
    _table(doc,
           ["Métrica", "Valor", "Descripción"],
           [
               ["TAM", "$6.78B USD",
                "Mercado global de gestión documental e inteligencia empresarial (2026)"],
               ["SAM", "$800M USD",
                "Mercado de Latinoamérica (México, Centroamérica, Colombia, Perú, Chile)"],
               ["SOM", "$45M USD",
                "Mercado alcanzable inicial: Panamá + Costa Rica + Colombia en 3 años"],
           ],
           col_widths=[2.0, 3.0, 11.0])

    _heading(doc, "8.2 Segmentación por Fases", level=2)
    for phase, desc in [
        ("Fase 1 — Panamá (Meses 1-12): ",
         "Despachos legales, consultoras, universidades (UTP, UP, USMA), bancos (compliance), constructoras."),
        ("Fase 2 — Centroamérica (Año 2): ",
         "Costa Rica (sector tecnológico y servicios) y Colombia (mercado grande)."),
        ("Fase 3 — Latinoamérica (Año 3+): ",
         "México (el más grande de LA), Perú, Chile, Ecuador."),
    ]:
        _bullet(doc, desc, bold_prefix=phase)

    _heading(doc, "8.3 Buyer Persona — Cliente Ideal", level=2)
    _table(doc,
           ["Atributo", "Descripción"],
           [
               ["Nombre", "Carlos Méndez"],
               ["Edad", "42 años"],
               ["Rol", "Gerente General de constructora mediana (50 empleados)"],
               ["Ubicación", "Panamá City, Panamá"],
               ["Dolor principal", "Pierde 5 horas/semana buscando cláusulas de contratos y respondiendo "
                                   "preguntas de empleados sobre políticas internas"],
               ["Presupuesto mensual para software", "$200 – $500 USD"],
               ["Decisión de compra", "Recomendación de colega + demo gratuita de 14 días"],
           ],
           col_widths=[5.0, 11.0])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 10. COMPETENCIAS
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "9. Competencias", level=1)

    _heading(doc, "9.1 Análisis Competitivo", level=2)
    _table(doc,
           ["Competidor", "Precio", "Fortaleza", "Debilidad vs. DocuMind AI"],
           [
               ["ChatGPT Enterprise", "$30/usuario/mes",
                "Conocimiento general potente",
                "No conoce docs internos. Alucinaciones. Costoso por usuario."],
               ["Microsoft Copilot", "$30/usuario/mes",
                "Integrado con Word/Excel/Teams",
                "Requiere ecosistema Microsoft completo. Precio alto."],
               ["Glean", "$$$ (enterprise)", "Búsqueda potente en grandes empresas",
                "Solo para empresas 500+ empleados. Precio prohibitivo para MIPYMES."],
               ["Notion AI", "$10/usuario/mes",
                "Buena para notas y wikis",
                "No diseñado para PDFs legales complejos. No cita fuentes exactas."],
               ["AskYourPDF", "$15/mes",
                "Simple y fácil de usar",
                "Solo 1 PDF a la vez. No multi-documento. No enterprise."],
               ["DocuMind AI", "$49 – $999/mes",
                "Español nativo, citas exactas, on-premise opcional, precio latinoamericano",
                "Startup nueva, falta de marca reconocida (ventaja: sin deuda técnica)"],
           ],
           col_widths=[3.5, 2.5, 4.5, 5.5])

    _heading(doc, "9.2 Matriz de Características", level=2)
    _table(doc,
           ["Característica", "DocuMind", "ChatGPT", "Copilot", "Glean", "Notion AI"],
           [
               ["Citas exactas (archivo, página)", "✅ SÍ", "❌ NO", "⚠️ PARCIAL", "❌ NO", "❌ NO"],
               ["Multi-documento", "✅ SÍ", "❌ NO", "✅ SÍ", "✅ SÍ", "❌ NO"],
               ["Español latinoamericano", "✅ SÍ", "⚠️ GENÉRICO", "⚠️ GENÉRICO", "⚠️ GENÉRICO", "⚠️ GENÉRICO"],
               ["On-premise opcional", "✅ SÍ", "❌ NO", "❌ NO", "❌ NO", "❌ NO"],
               ["Precio para MIPYMES", "✅ $49-$199", "⚠️ $30/user", "⚠️ $30/user", "❌ $$$$", "⚠️ $10/user"],
               ["Stack open source", "✅ SÍ", "❌ NO", "❌ NO", "❌ NO", "❌ NO"],
               ["Sin dependencia Big Tech", "✅ SÍ", "❌ NO", "❌ NO", "❌ NO", "❌ NO"],
           ],
           col_widths=[5.5, 2.2, 2.2, 2.2, 2.2, 2.2])

    _heading(doc, "9.3 Ventaja Diferencial", level=2)
    _body(doc,
          "La trazabilidad exacta: ningún competidor muestra 'página 7, párrafo 3' con tanta precisión. "
          "Esto es decisivo para abogados (citar contratos en tribunales), auditores (rastrear políticas), "
          "y compliance officers (demostrar procedimientos seguidos).",
          color=VIOLET)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 11. MODELO DE NEGOCIO
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "10. Modelo de Negocio", level=1)

    _heading(doc, "10.1 Planes de Suscripción", level=2)
    _table(doc,
           ["Plan", "Precio Mensual", "Usuarios", "Almacenamiento", "Características"],
           [
               ["Free", "$0", "1", "500 MB",
                "3 documentos máx., 20 consultas/mes, sin OCR. Gancho de adquisición."],
               ["Starter", "$49 USD", "5", "5 GB",
                "Documentos ilimitados, OCR activo, chat RAG, búsqueda semántica, soporte email."],
               ["Business", "$199 USD", "50", "100 GB",
                "Todo Starter + multi-workspace, analytics avanzados, API access, soporte prioritario."],
               ["Enterprise", "$999 USD", "Ilimitados", "Ilimitado",
                "Todo Business + on-premise, integración SAP/SharePoint, SSO/SAML, CSM dedicado, SLA 99.9%."],
           ],
           col_widths=[2.5, 3.0, 2.5, 3.5, 4.5])

    _heading(doc, "10.2 ARPU y Margen por Plan", level=2)
    _table(doc,
           ["Plan", "ARPU Mensual", "ARPU Anual", "Margen Bruto"],
           [
               ["Starter", "$49", "$588", "75%"],
               ["Business", "$199", "$2,388", "70%"],
               ["Enterprise", "$999", "$11,988", "65%"],
           ],
           col_widths=[3.5, 3.5, 3.5, 5.5])

    _heading(doc, "10.3 Estrategia de Precios", level=2)
    for item in [
        "Pago anual: 2 meses gratis (≈17% de descuento).",
        "Early adopters: 50% de descuento el primer año a cambio de testimonio público.",
        "Organizaciones sin fines de lucro: 30% de descuento permanente.",
        "Starter ($49) accesible para microempresas panameñas — el punto de entrada ideal.",
    ]:
        _bullet(doc, item)

    _heading(doc, "10.4 Estrategia Freemium y Conversión", level=2)
    _body(doc,
          "El Plan Free es el gancho de adquisición más poderoso de DocuMind AI. "
          "Permite hasta 3 documentos y 20 consultas mensuales sin costo. "
          "Cuando el usuario alcanza el límite (visiblemente en el dashboard), "
          "recibe un bloqueo amigable con CTA directo a Stripe Checkout.")
    for item in [
        "Conversión esperada Free → Starter: 8-12% (benchmarks SaaS B2B similares).",
        "Sin tarjeta de crédito al registrarse → reduce fricción de entrada.",
        "Upgrade en menos de 2 clics: UpgradeDialog → Stripe → plan activo de inmediato.",
    ]:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 12. MARKETING
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "11. Marketing", level=1)

    _heading(doc, "11.1 Validación (Fase 0 — Mes 0-1)", level=2)
    _body(doc,
          "Confirmar que empresas reales pagarían por DocuMind antes de invertir en desarrollo completo.")
    for item in [
        "Landing page con formulario de waitlist. Meta: 100 correos en 30 días.",
        "Oferta: 'Acceso anticipado gratuito para las primeras 50 empresas.'",
        "Métrica de éxito: 20+ correos de empresas reales (no amigos/familia) validan el problema.",
        "Canales: LinkedIn (posts sobre productividad en Panamá), grupos Facebook, WhatsApp Business.",
    ]:
        _bullet(doc, item)

    _heading(doc, "11.2 Adquisición de Clientes (Fase 1 — Meses 1-6)", level=2)
    _body(doc, "Canales pagados (bajo presupuesto):", bold=True)
    for item in [
        "Google Ads: $200/mes en keywords 'software gestión documental Panamá'.",
        "LinkedIn Ads: $100/mes targeting gerentes de RRHH y legales en Panamá.",
    ]:
        _bullet(doc, item)
    _body(doc, "Canales orgánicos (gratuitos):", bold=True)
    for item in [
        "Blog SEO: 'Cómo reducir 50% el tiempo de búsqueda de documentos en tu empresa.'",
        "YouTube: Demos de 2 minutos mostrando el producto en acción.",
        "LinkedIn personal: 3 posts semanales sobre productividad y tecnología.",
    ]:
        _bullet(doc, item)

    _heading(doc, "11.3 Alianzas Estratégicas (Fase 2 — Meses 6-12)", level=2)
    for partner, desc in [
        ("AMPYME: ", "Alianza para ofrecer DocuMind a emprendedores formalizados. Acceso a la red de MiPymes."),
        ("Cámara de Comercio: ", "Demo en eventos de networking. Descuento especial para socios. Testimonio institucional."),
        ("Universidades (UTP, UP, USMA): ", "Licencias gratuitas para uso académico + casos de estudio en clases."),
        ("Despachos de abogados: ", "Programa early adopters con 50% de descuento + casos de éxito publicados."),
    ]:
        _bullet(doc, desc, bold_prefix=partner)

    _heading(doc, "11.4 Escalamiento (Fase 3 — Año 2+)", level=2)
    for item in [
        "Expansión geográfica: Costa Rica → Colombia → México.",
        "API pública para desarrolladores.",
        "Marketplace de integraciones: Google Drive, Dropbox, OneDrive, SAP.",
        "White-label para consultoras que quieran ofrecer DocuMind bajo su marca.",
        "Resellers locales en cada país + inside sales.",
    ]:
        _bullet(doc, item)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 13. OBJETIVOS DE CRECIMIENTO
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "12. Objetivos de Crecimiento", level=1)

    _heading(doc, "12.1 Corto Plazo (0-6 meses)", level=2)
    _table(doc,
           ["Objetivo", "Meta", "KPI"],
           [
               ["Lanzar MVP funcional", "100% completo", "Demo funcional sin errores"],
               ["Conseguir empresas piloto", "10 empresas", "Feedback cualitativo recopilado"],
               ["Validar modelo de precios", "5 pagos reales", "Revenue > $0 confirmado"],
               ["Registrar marca DocuMind", "100% completo", "Registro en DPI Panamá"],
               ["Constituir empresa (SA Panamá)", "100% completo", "SA de Panamá activa"],
               ["Conseguir waitlist", "500 suscriptores", "50% de empresas reales"],
           ],
           col_widths=[6.0, 3.5, 6.5])

    _heading(doc, "12.2 Mediano Plazo (6-18 meses)", level=2)
    _table(doc,
           ["Objetivo", "Meta", "KPI"],
           [
               ["Alcanzar clientes pagos", "100 clientes", "MRR > $10,000"],
               ["Lanzar versión 2.0", "100% completo", "Multi-tenancy real + API pública"],
               ["Expandir geográficamente", "2 países nuevos", "Costa Rica y Colombia operativos"],
               ["Equipo completo", "5 personas", "2 devs, 1 ventas, 1 soporte, 1 CEO"],
               ["Levantar ronda semilla", "$50,000 – $100,000", "Term sheet firmado"],
               ["Alianzas estratégicas", "3 activas", "AMPYME, Cámara de Comercio, 1 universidad"],
           ],
           col_widths=[6.0, 3.5, 6.5])

    _heading(doc, "12.3 Largo Plazo (18-36 meses)", level=2)
    _table(doc,
           ["Objetivo", "Meta", "KPI"],
           [
               ["Clientes activos", "1,000+", "ARR > $1M"],
               ["Cobertura geográfica", "5 países", "Panamá, CR, Colombia, México, Perú"],
               ["Módulos avanzados", "100% activos", "Compliance + Insights funcionando"],
               ["Integraciones gubernamentales", "3 activas", "DGI, CSS, Municipio de Panamá"],
               ["Expansión a hispanos en USA", "1 mercado", "Miami / Texas operativo"],
               ["Opciones de salida", "2 opciones", "Adquisición por Box/Dropbox o IPO local"],
           ],
           col_widths=[6.0, 3.5, 6.5])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 14. FINANZAS
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "13. Finanzas — Proyección a 3 Años", level=1)

    _heading(doc, "13.1 Supuestos Clave", level=2)
    _table(doc,
           ["Supuesto", "Valor"],
           [
               ["Precio promedio ponderado por cliente", "$120/mes"],
               ["Churn anual", "15%"],
               ["Crecimiento mensual clientes (meses 1-6)", "2-3 clientes/mes"],
               ["Crecimiento mensual clientes (meses 7-12)", "5 clientes/mes"],
               ["Crecimiento mensual clientes (Año 2)", "8-10 clientes/mes"],
               ["Crecimiento mensual clientes (Año 3)", "15-20 clientes/mes"],
               ["CAC (Costo de Adquisición por Cliente)", "$50/cliente"],
               ["LTV (Lifetime Value)", "$1,440 ($120 × 12 × 85% retención / 15% churn)"],
               ["Ratio LTV/CAC", "28.8 (excelente; >3 es saludable)"],
           ],
           col_widths=[8.0, 8.0])

    _heading(doc, "13.2 Proyección Financiera (3 Años)", level=2)
    _table(doc,
           ["Métrica", "Año 1", "Año 2", "Año 3"],
           [
               ["Clientes nuevos", "20", "80", "300"],
               ["Clientes acumulados (fin de año)", "20", "85", "340"],
               ["Clientes recurrentes (fin de año)", "17", "72", "289"],
               ["Churn (clientes perdidos)", "3", "13", "51"],
               ["MRR (Monthly Recurring Revenue)", "$2,040", "$8,640", "$34,680"],
               ["ARR (Annual Recurring Revenue)", "$12,000", "$60,000", "$300,000"],
               ["Costos operativos anuales", "$8,000", "$35,000", "$120,000"],
               ["Utilidad neta anual", "$4,000", "$25,000", "$180,000"],
               ["Margen neto", "33%", "42%", "60%"],
               ["Break-even", "Mes 10", "—", "—"],
           ],
           col_widths=[7.0, 3.0, 3.0, 3.0])

    _heading(doc, "13.3 Desglose de Costos Operativos", level=2)
    _body(doc, "Año 1 — $8,000 total:", bold=True)
    _table(doc,
           ["Concepto", "Mensual", "Anual"],
           [
               ["Servidor cloud (GPU A4000)", "$209", "$2,508"],
               ["Marketing digital (LinkedIn, Google Ads)", "$200", "$2,400"],
               ["Herramientas (GitHub, Figma, Notion)", "$50", "$600"],
               ["Legal y contable", "$200", "$2,400"],
               ["Internet y servicios básicos", "$8", "$92"],
               ["TOTAL", "$667", "$8,000"],
           ],
           col_widths=[9.0, 3.5, 3.5])

    _body(doc, "Año 2 — $35,000 total:", bold=True)
    _table(doc,
           ["Concepto", "Mensual", "Anual"],
           [
               ["Servidores cloud (escalado)", "$800", "$9,600"],
               ["Marketing (aumentado)", "$500", "$6,000"],
               ["Salarios (2 devs part-time)", "$1,000", "$12,000"],
               ["Herramientas y software", "$100", "$1,200"],
               ["Legal, contable, impuestos", "$300", "$3,600"],
               ["Viajes y eventos", "$200", "$2,400"],
               ["TOTAL", "$2,900", "$35,000"],
           ],
           col_widths=[9.0, 3.5, 3.5])

    _body(doc, "Año 3 — $120,000 total:", bold=True)
    _table(doc,
           ["Concepto", "Mensual", "Anual"],
           [
               ["Infraestructura cloud (escalado)", "$3,000", "$36,000"],
               ["Marketing y ventas", "$2,000", "$24,000"],
               ["Salarios (5 personas full-time)", "$4,000", "$48,000"],
               ["Herramientas y software", "$300", "$3,600"],
               ["Legal, contable, impuestos", "$500", "$6,000"],
               ["Oficina y operaciones", "$200", "$2,400"],
               ["TOTAL", "$10,000", "$120,000"],
           ],
           col_widths=[9.0, 3.5, 3.5])

    _heading(doc, "13.4 Punto de Equilibrio (Break-Even)", level=2)
    _table(doc,
           ["Parámetro", "Valor"],
           [
               ["Inversión inicial", "$4,100 USD"],
               ["Costos fijos mensuales (Año 1)", "$667 USD/mes"],
               ["Ingreso promedio por cliente", "$120 USD/mes"],
               ["Clientes necesarios para break-even", "6 clientes"],
               ["Mes estimado de break-even", "Mes 10"],
           ],
           col_widths=[8.0, 8.0])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 15. CONCLUSIÓN
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "14. Conclusión", level=1)

    _heading(doc, "14.1 Resumen del Proyecto", level=2)
    _body(doc,
          "DocuMind AI es una startup tecnológica panameña que resuelve un problema real y "
          "cuantificable: las empresas pierden miles de horas anuales buscando información en "
          "documentos internos. Mediante inteligencia artificial (RAG + embeddings + LLMs open source), "
          "DocuMind convierte documentos estáticos en conocimiento accionable, citando la fuente exacta "
          "con precisión de archivo, página y párrafo.")
    _body(doc,
          "El proyecto se fundamenta en conceptos clave de la materia de Innovación y Emprendimiento "
          "Empresarial: innovación tecnológica, modelo de negocio SaaS escalable, transformación digital "
          "de procesos empresariales, liderazgo innovador y trabajo en equipo.")

    _heading(doc, "14.2 Fortalezas del Proyecto", level=2)
    for item in [
        ("Tecnología probada: ", "El stack RAG es tecnología de vanguardia en 2026, usada por Microsoft, Google y startups de Y Combinator."),
        ("Mercado validado: ", "El 95% de empresas en Panamá son MIPYMES sin herramientas para gestión documental inteligente."),
        ("Equipo capacitado: ", "Desarrolladores con experiencia en IA, pipelines de datos y desarrollo web moderno."),
        ("Costo de entrada bajo: ", "$4,100 de inversión inicial, operable con recursos propios (bootstrapping)."),
        ("Escalabilidad: ", "Modelo SaaS con márgenes crecientes (33% Año 1 → 60% Año 3)."),
    ]:
        _bullet(doc, item[1], bold_prefix=item[0])

    _heading(doc, "14.3 Próximos Pasos", level=2)
    for item in [
        "Semana 1: Completar MVP funcional para presentación semestral.",
        "Mes 1: Lanzar landing page y capturar 100 emails de waitlist.",
        "Meses 2-3: Onboarding de 10 empresas piloto en Panamá.",
        "Meses 4-6: Iterar producto con feedback, lanzar versión 1.0 pública.",
        "Meses 7-12: Escalar a 100 clientes pagos, expandir equipo.",
        "Año 2: Expansión a Costa Rica y Colombia, levantar ronda semilla $50K-$100K.",
    ]:
        _bullet(doc, item)

    _heading(doc, "14.4 Llamado a la Acción", level=2)
    _body(doc,
          "DocuMind AI no es solo un chatbot de PDFs. Es el Sistema Operativo del Conocimiento "
          "Empresarial que transformará cómo las empresas latinoamericanas gestionan su información. "
          "Con una inversión inicial de $4,100, un mercado de $45M alcanzable, y un equipo "
          "comprometido, estamos listos para convertir el caos documental en productividad.",
          bold=True, color=NAVY)
    _body(doc,
          '"Innovar no es solo crear algo nuevo, es crear algo que valga la pena." — DocuMind AI',
          italic=True, color=VIOLET)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 16. COSTOS DE LA STARTUP — ACTIVOS
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "15. Costos de la Startup — Activos", level=1)

    _heading(doc, "15.1 Inversión Inicial Requerida", level=2)
    _table(doc,
           ["Concepto", "Monto (USD)", "Descripción"],
           [
               ["Laptop de desarrollo (GPU dedicada)", "$1,200",
                "Laptop con 16GB RAM mínimo para desarrollo y demo"],
               ["Dominio + hosting (1 año)", "$100",
                "documind.ai + hosting básico en Vercel/Railway"],
               ["Branding y diseño", "$300",
                "Logo profesional, mockups, paleta de colores, iconografía"],
               ["Registro de empresa (SA Panamá)", "$500",
                "Constitución en Registro Público, patente comercial"],
               ["Capital de trabajo (3 meses)", "$2,000",
                "Cubre costos operativos mientras se consiguen primeros clientes"],
               ["TOTAL INVERSIÓN INICIAL", "$4,100", ""],
           ],
           col_widths=[5.0, 2.5, 8.5])

    _heading(doc, "15.2 Activos Fijos", level=2)
    _table(doc,
           ["Activo", "Valor Inicial", "Vida Útil", "Depreciación Anual"],
           [
               ["Equipos informáticos (laptops × 4)", "$4,800", "3 años", "$1,600"],
               ["Mobiliario de oficina (escritorios, sillas)", "$600", "5 años", "$120"],
               ["Servidor dedicado (Año 2)", "$2,500", "3 años", "$833"],
               ["TOTAL ACTIVOS FIJOS", "$7,900", "", "$2,553"],
           ],
           col_widths=[6.0, 3.0, 2.5, 4.5])

    _heading(doc, "15.3 Activos Intangibles", level=2)
    _table(doc,
           ["Activo", "Valor", "Descripción"],
           [
               ["Software propietario (DocuMind AI)", "$15,000",
                "Valor estimado del código fuente y algoritmos"],
               ["Marca registrada", "$500",
                "Registro de marca en DPI Panamá"],
               ["Base de conocimiento (embeddings)", "$2,000",
                "Valor de los datos y vectores acumulados"],
               ["Relaciones comerciales", "$1,000",
                "Red de contactos, aliados, clientes potenciales"],
               ["TOTAL ACTIVOS INTANGIBLES", "$18,500", ""],
           ],
           col_widths=[5.5, 2.5, 8.0])

    _heading(doc, "15.4 Capital de Trabajo", level=2)
    _table(doc,
           ["Concepto", "Mes 1-3", "Mes 4-6", "Mes 7-12"],
           [
               ["Efectivo disponible", "$2,000", "$1,500", "$3,000"],
               ["Cuentas por cobrar", "$0", "$500", "$2,000"],
               ["Inventario (licencias prepagadas)", "$0", "$0", "$500"],
               ["Cuentas por pagar", "$500", "$800", "$1,200"],
               ["Capital de trabajo neto", "$1,500", "$1,200", "$4,300"],
           ],
           col_widths=[5.5, 3.0, 3.0, 4.5])

    _heading(doc, "15.5 Desglose de Costos de Puesta en Marcha", level=2)
    _table(doc,
           ["Categoría", "Costo (USD)", "% del Total"],
           [
               ["Tecnología y equipos", "$1,500", "37%"],
               ["Legal y administrativo", "$800", "20%"],
               ["Marketing y ventas inicial", "$600", "15%"],
               ["Capital de trabajo", "$1,000", "24%"],
               ["Contingencia (10%)", "$200", "5%"],
               ["TOTAL", "$4,100", "100%"],
           ],
           col_widths=[7.0, 4.0, 5.0])

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # REFERENCIAS
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "Referencias", level=1)
    refs = [
        "AMPYME — Autoridad de la Micro, Pequeña y Mediana Empresa de Panamá. Datos estadísticos de MIPYMES 2026.",
        "La Prensa Panamá. \"MiPymes: motor clave del 70% de la economía panameña.\" 28 de abril de 2026.",
        "Ministerio de Comercio e Industrias de Panamá. Ley No. 16 de 2008 (Ley de MIPYMES).",
        "Gartner Research. \"Enterprise Document Management Market Analysis 2026.\"",
        "LangChain Documentation. \"Retrieval-Augmented Generation (RAG) Best Practices.\" 2026.",
        "Groq Cloud. \"API Documentation — Llama 3.1 Models.\" 2026.",
        "ChromaDB. \"Vector Database for AI Applications.\" 2026.",
        "Clerk. \"Next.js Authentication Guide.\" 2026.",
        "Stripe. \"Subscriptions and Checkout.\" 2026.",
        "Tesseract OCR. \"Documentation and Language Models.\" 2026.",
        "Materials de clase: Glosario de Innovación y Emprendimiento Empresarial, UTP 2026.",
        "Materials de clase: Importancia de las MIPYMES en Panamá, UTP 2026.",
        "Materials de clase: Módulo I — Filosofía del Emprendedor, UTP 2026.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="List Number")
        _paragraph_spacing(p, before=0, after=30)
        p.add_run(ref).font.size = Pt(10)

    _page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ANEXOS
    # ─────────────────────────────────────────────────────────────────────────
    _heading(doc, "Anexos", level=1)

    _heading(doc, "Anexo A — Documentos Demo Completos", level=2)
    _body(doc,
          "Los 7 documentos de Constructora Nova S.A. se generan automáticamente con el script "
          "backend/scripts/generate_demo_docs.py y se indexan en el startup de FastAPI. "
          "Ver Sección 6.2 para el listado completo con contenido clave y páginas de referencia.")

    _heading(doc, "Anexo B — Preguntas de Demo con Respuestas (12 preguntas)", level=2)
    _table(doc,
           ["#", "Pregunta", "Respuesta Esperada", "Documento", "Página"],
           [
               ["1", "¿Cuántos días de vacaciones tiene un empleado con 2 años?",
                "17 días hábiles", "Manual_RRHH_Nova.pdf", "4"],
               ["2", "¿Qué pasa si renuncia sin preaviso?",
                "Pérdida de beneficios", "Reglamento_Interno_Nova.pdf", "3"],
               ["3", "¿Plazo del arrendamiento?",
                "24 meses renovables", "Contrato_Arrendamiento_Oficina.pdf", "2"],
               ["4", "¿Preaviso de renuncia?",
                "30 días calendario", "Contrato_Laboral_Nova_2026.pdf", "3"],
               ["5", "¿Quién aprueba compra de $3,000?",
                "Gerente General + Gerente de Finanzas", "Politica_Compras_Proveedores.pdf", "4"],
               ["6", "¿Horario laboral?",
                "Lunes a viernes 8am-5pm", "Manual_RRHH_Nova.pdf", "2"],
               ["7", "¿Cobertura seguro médico emergencias?",
                "100% sin deducible", "Contrato_Laboral_Nova_2026.pdf", "4"],
               ["8", "¿Causales de despido inmediato?",
                "Robo, violencia, secretos, ausencia", "Manual_RRHH_Nova.pdf", "6"],
               ["9", "¿Canon mensual de alquiler?",
                "$3,500.00 USD", "Contrato_Arrendamiento_Oficina.pdf", "3"],
               ["10", "¿Cotizaciones para compras >$500?",
                "Al menos 3 cotizaciones", "Politica_Compras_Proveedores.pdf", "3"],
               ["11", "¿Qué dice el acta sobre el presupuesto?",
                "Presupuesto anual 2026: $3,200,000 aprobado", "Acta_Reunion_Directiva_Nova.jpg", "1"],
               ["12", "¿Cuánto dice la factura que se debe pagar?",
                "$8,160.00 USD", "Factura_Servicios_Enero_2026.pdf", "1"],
           ],
           col_widths=[0.7, 5.5, 4.0, 4.5, 1.3])

    _heading(doc, "Anexo C — Business Model Canvas", level=2)
    _table(doc,
           ["Bloque", "Contenido"],
           [
               ["Propuesta de valor", "Sistema Operativo del Conocimiento Empresarial con citas exactas (archivo, página, párrafo)"],
               ["Segmentos de clientes", "Despachos legales, consultoras, constructoras, universidades, bancos (compliance)"],
               ["Canales", "Landing page, LinkedIn, alianzas AMPYME, Cámara de Comercio, referidos"],
               ["Relaciones", "Soporte email (Starter), CSM dedicado (Enterprise), comunidad online"],
               ["Fuentes de ingreso", "Suscripción SaaS mensual/anual (Free → Starter $49 → Business $199 → Enterprise $999)"],
               ["Recursos clave", "Código propietario, embeddings locales, marca registrada, equipo técnico"],
               ["Actividades clave", "Desarrollo de IA/RAG, ventas B2B, soporte al cliente, marketing digital"],
               ["Alianzas clave", "AMPYME, Cámara de Comercio, universidades, consultoras tecnológicas"],
               ["Estructura de costos", "Infraestructura cloud, marketing digital, salarios, legal/contable"],
           ],
           col_widths=[4.0, 12.0])

    _heading(doc, "Anexo D — Glosario de Términos", level=2)
    glosario = [
        ("RAG (Retrieval-Augmented Generation): ",
         "Técnica de IA que combina recuperación de información (búsqueda vectorial) con generación de texto (LLM) para producir respuestas fundamentadas en documentos reales."),
        ("Embedding: ",
         "Representación numérica (vector) del significado semántico de un texto, que permite comparar similitudes entre frases."),
        ("ChromaDB: ",
         "Base de datos vectorial open-source que almacena embeddings y permite búsquedas semánticas eficientes."),
        ("Chunking: ",
         "División de documentos en fragmentos más pequeños (chunks) para indexarlos y recuperarlos de forma granular."),
        ("LLM (Large Language Model): ",
         "Modelo de lenguaje a gran escala entrenado con millones de textos. DocuMind usa Llama 3.1 vía Groq API."),
        ("OCR (Optical Character Recognition): ",
         "Tecnología que extrae texto de imágenes o documentos escaneados. DocuMind usa Tesseract OCR."),
        ("SaaS (Software as a Service): ",
         "Modelo de negocio donde el software se ofrece como servicio vía suscripción mensual."),
        ("MRR / ARR: ",
         "Monthly Recurring Revenue / Annual Recurring Revenue. Métricas clave de ingresos para startups SaaS."),
        ("CAC / LTV: ",
         "Customer Acquisition Cost / Lifetime Value. Ratio LTV/CAC > 3 indica modelo de negocio saludable."),
        ("Gating: ",
         "Bloqueo de funcionalidades según el plan del usuario, implementado en el proxy de Next.js."),
    ]
    for term, definition in glosario:
        _bullet(doc, definition, bold_prefix=term)

    # ─────────────────────────────────────────────────────────────────────────
    # Guardar
    # ─────────────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"[OK] Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build_document()
